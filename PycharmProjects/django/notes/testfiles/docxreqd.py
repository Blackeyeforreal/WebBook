import docx.table
import docx
import docx2python
import re
import sys
import os
import comtypes.client
import json

class docx1():
    def __init__(self):
        self.play_notes = docx.Document('D:/project/django/Notes from The Other Einstein.docx' )
        self.tables= self.play_notes.tables
        self.pdf_words= docx.Document()
        self.pdf_quotes = docx.Document()
        self.t = self.pdf_words.add_table(1, 3)
        self.f = open("../notes/dictionary.json")
        self.t.autofit = False
        self.data = json.load(self.f)
        self.count = 0

    def pdftodocx(self,in_file,out_file):
        wdFormatPDF = 17

        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

    def automation(self):
        for table in self.tables:
              for cell in table.cell(0,0).tables:
                    k = cell.cell(0,1).text
                    x = re.search(r"[A-Za-z]\w* [\d]{1,2}, [\d]{4}", k)
                    first,secord= x.span()
                    y = re.search(r"\w ", k[:first])
                    self.t.autofit= True

                    if y== None:
                        self.count +=1 #detecting words

                        self.t.cell(self.count -1,0).add_paragraph(f"{self.count}")
                        self.t.cell(self.count -1,0).width= 56
                        self.t.cell(self.count -1,1).add_paragraph(k[:first-2])
                        self.t.add_row()




                        try:
                            self.t.cell(self.count -1 ,2).add_paragraph(self.data[k[:first-2]])
                        except:
                            self.t.cell(self.count -1 ,2).add_paragraph("couldn,t find")



                    else:
                        s,e=y.span()
                        print("YOUR STRING")
                        self.pdf_quotes.add_paragraph(k[:first])

                        print(k[:first])
        for cell in self.t.columns[0].cells:
            cell.width = 45


        #pdf_quotes.save("pdf_quotes.docx")
        self.pdf_words.save("pdf_words.docx")
       # self.pdftodocx(r"C:\Users\devang srivastava\PycharmProjects\django\notes\testfiles\pdf_words.docx",r"C:\Users\devang srivastava\PycharmProjects\django\notes\testfiles\pdf_words.pdf")

        self.pdftodocx(r"C:\Users\devang srivastava\PycharmProjects\django\notes\media\docx\jjj\pdf_words.docx",r"C:\Users\devang srivastava\PycharmProjects\django\notes\testfiles\pdf_words.pdf")

    def oncalledit(self,num):
        para = self.t.cell(num,1).text
        return  para

k =docx1()
k.automation()



